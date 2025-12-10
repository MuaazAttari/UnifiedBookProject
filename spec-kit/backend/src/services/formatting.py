import asyncio
import tempfile
import os
from typing import Dict, List
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches
from sqlalchemy.orm import Session

from ..models.textbook import Textbook
from ..models.chapter import Chapter
from ..models.section import Section


class FormattingService:
    """
    Service for formatting and exporting textbooks in different formats
    """

    def __init__(self):
        pass

    def _get_textbook_data(self, db: Session, textbook_id: str) -> Dict:
        """
        Retrieve textbook data with all chapters and sections
        """
        textbook = db.query(Textbook).filter(Textbook.id == textbook_id).first()
        if not textbook:
            raise ValueError(f"Textbook with id {textbook_id} not found")

        chapters = db.query(Chapter).filter(Chapter.textbook_id == textbook_id).order_by(Chapter.position).all()

        textbook_data = {
            "id": textbook.id,
            "title": textbook.title,
            "subject": textbook.subject,
            "educational_level": textbook.educational_level,
            "chapters": []
        }

        for chapter in chapters:
            sections = db.query(Section).filter(Section.chapter_id == chapter.id).order_by(Section.position).all()

            chapter_data = {
                "id": chapter.id,
                "title": chapter.title,
                "position": chapter.position,
                "status": chapter.status,
                "sections": []
            }

            for section in sections:
                section_data = {
                    "id": section.id,
                    "title": section.title,
                    "position": section.position,
                    "type": section.type,
                    "content": section.content
                }
                chapter_data["sections"].append(section_data)

            textbook_data["chapters"].append(chapter_data)

        return textbook_data

    def _generate_html_content(self, textbook_data: Dict, include_solutions: bool = False) -> str:
        """
        Generate HTML content for the textbook
        """
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{textbook_data['title']}</title>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #333; text-align: center; }}
                h2 {{ color: #555; border-bottom: 2px solid #007acc; }}
                h3 {{ color: #666; }}
                .chapter {{ margin-bottom: 30px; }}
                .section {{ margin-bottom: 20px; }}
                .exercise {{ background-color: #f9f9f9; padding: 10px; border-left: 4px solid #007acc; }}
                .summary {{ background-color: #e8f4f8; padding: 10px; border-left: 4px solid #28a745; }}
                .key-point {{ background-color: #fff3cd; padding: 10px; border-left: 4px solid #ffc107; }}
            </style>
        </head>
        <body>
            <h1>{textbook_data['title']}</h1>
            <h3>Subject: {textbook_data['subject']}</h3>
            <h3>Educational Level: {textbook_data['educational_level']}</h3>
        """

        for chapter in textbook_data['chapters']:
            html_content += f"""
            <div class="chapter">
                <h2>Chapter {chapter['position']}: {chapter['title']}</h2>
            """

            for section in chapter['sections']:
                section_class = ""
                if section['type'] == 'EXERCISE':
                    section_class = "exercise"
                elif section['type'] == 'SUMMARY':
                    section_class = "summary"
                elif section['type'] == 'KEY_POINT':
                    section_class = "key-point"

                html_content += f"""
                <div class="section {section_class}">
                    <h3>{section['title']}</h3>
                    <p>{section['content']}</p>
                </div>
                """

            html_content += "</div>"

        html_content += """
        </body>
        </html>
        """

        return html_content

    def export_to_pdf(self, db: Session, textbook_id: str, include_solutions: bool = False) -> str:
        """
        Export textbook to PDF format
        """
        textbook_data = self._get_textbook_data(db, textbook_id)
        html_content = self._generate_html_content(textbook_data, include_solutions)

        # Create a temporary file for the PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
            # Convert HTML to PDF using WeasyPrint
            html_doc = HTML(string=html_content)
            html_doc.write_pdf(temp_pdf.name)
            return temp_pdf.name

    def export_to_docx(self, db: Session, textbook_id: str, include_solutions: bool = False) -> str:
        """
        Export textbook to DOCX format
        """
        textbook_data = self._get_textbook_data(db, textbook_id)

        # Create a new document
        doc = Document()

        # Add title page
        doc.add_heading(textbook_data['title'], 0)
        doc.add_paragraph(f"Subject: {textbook_data['subject']}")
        doc.add_paragraph(f"Educational Level: {textbook_data['educational_level']}")
        doc.add_page_break()

        # Add chapters and sections
        for chapter in textbook_data['chapters']:
            # Add chapter heading
            doc.add_heading(f"Chapter {chapter['position']}: {chapter['title']}", level=1)

            for section in chapter['sections']:
                # Add section title
                doc.add_heading(section['title'], level=2)

                # Add section content
                # Split content into paragraphs by double newlines
                paragraphs = section['content'].split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())

                # Add some space between sections
                doc.add_paragraph()

            # Add space between chapters
            doc.add_paragraph()

        # Create a temporary file for the DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
            doc.save(temp_docx.name)
            return temp_docx.name

    def export_to_html(self, db: Session, textbook_id: str, include_solutions: bool = False) -> str:
        """
        Export textbook to HTML format
        """
        textbook_data = self._get_textbook_data(db, textbook_id)
        html_content = self._generate_html_content(textbook_data, include_solutions)

        # Create a temporary file for the HTML
        with tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as temp_html:
            temp_html.write(html_content)
            return temp_html.name


# Singleton instance
formatting_service = FormattingService()